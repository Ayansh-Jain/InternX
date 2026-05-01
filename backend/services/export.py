import io
import os
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_pdf_from_html(html_content: str) -> io.BytesIO:
    pdf_buffer = io.BytesIO()
    # Create PDF
    pisa_status = pisa.CreatePDF(
        src=html_content,
        dest=pdf_buffer
    )
    if pisa_status.err:
        raise Exception("Error generating PDF")
    
    pdf_buffer.seek(0)
    return pdf_buffer

def generate_docx_from_data(resume_data: dict) -> io.BytesIO:
    doc = Document()
    
    # Personal Info
    personal = resume_data.get('personal', {})
    name_paragraph = doc.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_paragraph.add_run(personal.get('fullName', '').upper())
    name_run.bold = True
    name_run.font.size = Pt(16)
    
    contact_parts = []
    if personal.get('email'): contact_parts.append(personal['email'])
    if personal.get('phone'): contact_parts.append(personal['phone'])
    if personal.get('location'): contact_parts.append(personal['location'])
    if personal.get('linkedIn'): contact_parts.append(personal['linkedIn'])
    
    contact_paragraph = doc.add_paragraph()
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_paragraph.add_run(" | ".join(contact_parts))
    
    # Experience
    if resume_data.get('experience'):
        doc.add_heading('Experience', level=1)
        for exp in resume_data['experience']:
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('role', '')}").bold = True
            p.add_run(f" | {exp.get('company', '')}")
            
            date_str = f" ({exp.get('startDate', '')} - {exp.get('endDate', 'Present')})"
            p.add_run(date_str)
            
            for bullet in exp.get('responsibilities', []):
                doc.add_paragraph(bullet, style='List Bullet')
                
    # Projects
    if resume_data.get('projects'):
        doc.add_heading('Projects', level=1)
        for proj in resume_data['projects']:
            p = doc.add_paragraph()
            p.add_run(f"{proj.get('title', '')}").bold = True
            if proj.get('description'):
                desc = proj.get('description')
                if isinstance(desc, list):
                    for d in desc:
                        doc.add_paragraph(d, style='List Bullet')
                elif isinstance(desc, str):
                    # split by period for bullets roughly
                    for d in [s.strip() for s in desc.split('.') if s.strip()]:
                        doc.add_paragraph(d, style='List Bullet')
                    
    # Education
    if resume_data.get('education'):
        doc.add_heading('Education', level=1)
        for edu in resume_data['education']:
            p = doc.add_paragraph()
            p.add_run(f"{edu.get('degree', '')}").bold = True
            p.add_run(f", {edu.get('college', '')}")
            if edu.get('endYear'):
                p.add_run(f" ({edu.get('startYear', '')} - {edu.get('endYear', '')})")

    # Skills
    skills = resume_data.get('skills', {})
    if skills:
        doc.add_heading('Skills', level=1)
        if skills.get('technical'):
            p = doc.add_paragraph()
            p.add_run("Technical: ").bold = True
            p.add_run(", ".join(skills['technical']))
        if skills.get('tools'):
            p = doc.add_paragraph()
            p.add_run("Tools: ").bold = True
            p.add_run(", ".join(skills['tools']))
        if skills.get('soft'):
            p = doc.add_paragraph()
            p.add_run("Soft Skills: ").bold = True
            p.add_run(", ".join(skills['soft']))

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer
